# -*- coding: utf-8 -*-
"""Emit docs/SRS_PEPSI_DOCX_REVISION_ITEM_BY_ITEM.md from docx body order."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
DOC = REPO / "docs" / "Software Requirement Specification Pepsi Cola PM Project.docx"
OUT = REPO / "docs" / "SRS_PEPSI_DOCX_REVISION_ITEM_BY_ITEM.md"

C1 = "`from customer/PM Application Requirement (Details).docx`"
C2 = "`from customer/PM Application Requirement (Details)Rev.1.docx`"
C3 = "`from customer/PM Application Requirement.docx`"
C4 = "`from customer/requirement_13_02_63 (003).docx`"
SET = f"{C1}, {C2}, {C3}, {C4}"


def action_table(tbl_idx: int, header: str) -> str:
    if tbl_idx == 0:
        return (
            "อัปเดต Version / Status / Date / Responsible; เพิ่มคำอธิบายว่าความต้องการหลักล็อกตามชุดลูกค้า 4 ไฟล์ (แบบ B) "
            "และเล่มนี้เป็น SRS เทคนิค/cross-ref"
        )
    if tbl_idx == 1:
        return f"เทียบทุกแถวกับ {C3} (stack) และ [`INFRASTRUCTURE.md`](INFRASTRUCTURE.md), [`INSTALL_SOP_TAILSCALE_DOCKER.md`](INSTALL_SOP_TAILSCALE_DOCKER.md)"
    if tbl_idx == 2:
        return (
            f"ซิงก์คอลัมน์รายละเอียดกับ {C2}+{C1} (ปฏิทิน M/W/D, Drag&Drop, สีงาน, Reason, TECO+ระฆัง); "
            f"หลังแก้ Word ให้อัปเดต [`SRS_TABLE_3_2.md`](SRS_TABLE_3_2.md)"
        )
    if 3 <= tbl_idx <= 10:
        return f"ตรวจทุกแถว 'รายละเอียด' ให้สอด UC ในหัวข้อด้านบน + {C2} (Precondition / Main / Exception / ข้อยกเว้น)"
    if tbl_idx == 11:
        return f"เทียบฟิลด์กับ {C2}+{C1} และ [`SAP_DATA_IMPORT_EXPORT_COLUMNS.md`](SAP_DATA_IMPORT_EXPORT_COLUMNS.md) (ชุด planning/IW37N)"
    if tbl_idx == 12:
        return f"เทียบฟิลด์กับ {C2}, `Hand held system.pptx`, [`SAP_DATA_IMPORT_EXPORT_COLUMNS.md`](SAP_DATA_IMPORT_EXPORT_COLUMNS.md)"
    if tbl_idx == 13:
        return f"เทียบฟิลด์กับ {C2} และคอลัมน์ GI/GR/Material ใน SAP_DATA_IMPORT_EXPORT_COLUMNS.md"
    if tbl_idx == 14:
        return f"เทียบฟิลด์รายงานกับ {C4} (dashboard/backlog) และนิยาม KPI ใน {C2}"
    if tbl_idx == 15:
        return f"เทียบกับ [`DATABASE_DESIGN_DRAFT.md`](DATABASE_DESIGN_DRAFT.md) และคอลัมน์จริงใน `from customer/SAP data` / {C3}"
    return f"ตรวจตารางกับ {SET}"


def action_paragraph(text: str) -> str:
    t = text.strip()
    # English parallel TOC lines (tab after section number, English title)
    if "\t" in t[:6] and re.match(r"^[1234]\.\d", t) and "วัตถุ" not in t and "บทที่" not in t:
        return "หัวข้อสารบัญภาษาอังกฤษ (คู่ขนาน): คงโครง; อัปเดตเลขหน้าเมื่อแก้เนื้อหาไทย"
    if len(t) < 80 and "แอปพลิเคชันบนเว็บ" in t:
        return "หน้าปก: อาจเพิ่มบรรทัดใต้ชื่อว่าความต้องการหลักอ้างชุด PM Application Requirement (from customer/)"
    if "Software Requirement Specification" in t and "ความต้องการทางด้านซอฟต์แวร์" in t:
        return "คำประกาศชื่อเอกสาร: เพิ่มคำว่า (Technical / aligned to customer Req*) ถ้าต้องการแยกบทบาทชัด"
    if t == "Document Story":
        return "คงหรือเปลี่ยนเป็น Document control ให้สอด Table 0"
    if t.startswith("san =") or t.startswith("chin =") or t.startswith("psit =") or t.startswith("khan = "):
        return "คำย่อชื่อทีม: ตรวจความเป็นปัจจุบัน (ไม่เกี่ยวลูกค้าโดยตรง)"
    if t == "Table of content":
        return "อัปเดตเลขหน้าใน TOC หลังแก้เนื้อหา"
    if t.startswith("Chapter "):
        return "หัว Chapter ภาษาอังกฤษ: คงโครง; แก้ถ้าเพิ่มบทใหม่"
    if "บทที่ 1" in t and "บทนำ" in t:
        return "หัวบทที่ 1 (ไทย): คงได้"
    if "1.1 วัตถุประสงค์" in t:
        return f"หัวข้อ 1.1: เปิด {C3} — ปรับย่อหน้าถัดไปให้วัตถุประสงค์ตรงกัน"
    if t.startswith("1. ") and "ความชัดเจนในการออกแบบ" in t:
        return f"ข้อวัตถุประสงค์ 1: เทียบ {C3}, {C4}"
    if t.startswith("2. ") and "แนวทางในการพัฒนา" in t:
        return f"ข้อวัตถุประสงค์ 2: เทียบ {C3} — ระบุว่าอ้างลูกค้าเป็นหลัก"
    if t.startswith("3. ") and "โครงสร้างพื้นฐาน" in t:
        return f"ข้อวัตถุประสงค์ 3: เทียบ {C3} (Docker/Tailscale) กับ INFRASTRUCTURE"
    if t.startswith("4. ") and "ประสิทธิภาพการจัดการข้อมูล" in t:
        return f"ข้อวัตถุประสงค์ 4: เทียบ {C3} (SAP reports)"
    if "1.2 ขอบเขตโครงการ" in t:
        return f"หัวข้อ 1.2: เตรียมเทียบข้อ 1–7 กับ {C3} และ {C4}"
    if t.startswith("โครงการนี้เกี่ยวข้องกับการพัฒนา"):
        return f"ย่อหน้านำขอบเขต: เทียบ {C3}"
    if "ระบบบริหารจัดการแผนงาน" in t and "Maintenance Scheduling" in t:
        return f"ข้อขอบเขต 1 (Scheduling): ซิงก์กับ {C2}+{C1} (ปฏิทิน, ZB01/02/05, DnD, Reason)"
    if "Hand-held" in t or "Handheld" in t and "ระบบปฏิบัติงาน" in t:
        return f"ข้อขอบเขต 2 (Handheld): เทียบ {C2} และ `Hand held system.pptx`"
    if "Asset & Resource" in t or "ทรัพยากรบุคคล" in t:
        return f"ข้อขอบเขต 3: เทียบ {C3} (Master data / FL)"
    if "แจ้งเตือนและสื่อสาร" in t or "Notification & Communication" in t:
        return f"ข้อขอบเขต 4: เทียบ {C3}/{C2} (Telegram, email)"
    if "Reporting & Dashboard" in t or "แดชบอร์ด" in t:
        return f"ข้อขอบเขต 5: เทียบ {C4}+{C2} (Backlog, KPI)"
    if "Infrastructure Mastery" in t or "โครงสร้างพื้นฐานและการจัดการข้อมูล" in t:
        return f"ข้อขอบเขต 6: เทียบ {C3} และ INSTALL_SOP / ภาคผนวก ค"
    if "Security & Access" in t or "Tailscale" in t:
        return f"ข้อขอบเขต 7 (VPN): เทียบ INSTALL_SOP และ {C3}"
    if "1.3 คำย่อ" in t:
        return "หัวข้อ 1.3: เตรียมเพิ่มนิยามจาก Rev.1"
    if t.startswith(("1. PM", "2. CM", "3. ZB02", "4. ZB01", "5. Order", "6. MntPlan", "7. CRTD", "8. REL", "9. TECO", "10. CLSD", "11. Functional", "12. Work Center", "13. Dashboard", "14. QR", "15. Docker", "16. Tailscale")):
        return f"นิยามรายการ: เทียบ {C2}+{C1}; เพิ่มคำว่า Reason code / สถานะสี / ระฆัง TECO ถ้ายังไม่ครบ"
    if "- TECO" in t or "การยืนยันปิดงานใน PM Application" in t:
        return f"คำอธิบาย TECO แยก SAP กับแอป: ต้องตรง {C2} (ระฆังเมื่อ TECO ใน SAP แต่ยังไม่ confirm ในแอป)"
    if "บทที่ 2" in t and "ภาพรวม" in t:
        return "หัวบทที่ 2: คงได้"
    if "2.1 มุมมอง" in t:
        return f"หัวข้อ 2.1: เปิด {C3} — ปรับย่อหน้าถัดไป"
    if "2.2 คุณสมบัติ" in t:
        return "หัวข้อ 2.2: คงได้"
    if t.startswith("จากการวิเคราะห์โครงสร้าง"):
        return f"คำนำ Feature: เทียบ {C3}"
    if t.startswith("Feature ") and "ระบบ" in t:
        return f"ข้อความ Feature: เทียบ {C2}+{C3} ทีละข้อ"
    if "2.3 กลุ่มผู้ใช้งาน" in t:
        return "หัวข้อ 2.3: คงได้"
    if "กลุ่มช่างเทคนิค" in t and "Technician" in t:
        return f"Planner/Technician บทบาท: เทียบ {C2} (หน้าที่ + สิทธิ์ลากงาน/Reason)"
    if "2.4 สภาพแวดล้อม" in t:
        return "หัวข้อ 2.4: คงได้"
    if "Network Environment" in t or "VPN Connection" in t or "Internet" in t:
        return "2.4 เครือข่าย: เทียบ INSTALL_SOP + ภาคผนวก ค"
    if "Server Side" in t:
        return "2.4 Server: เทียบ Table 1 + Docker + D:"
    if "Client Side" in t or "Web Browser" in t or "Handheld/Tablet" in t or "RAM" in t and "GB" in t:
        return f"2.4 Client: เทียบ {C2} (ปฏิทิน/Dashboard performance)"
    if "บทที่ 3" in t and "ฟังก์ชัน" in t:
        return "หัวบทที่ 3: คงได้"
    if "บทที่ 4" in t and "ความต้องการเฉพาะ" in t:
        return "หัวบทที่ 4: คงได้ — เนื้อหาด้านล่างเทียบทั้งบทกับ Rev.1"
    if "3.1" in t and "URS" in t:
        return "หัวข้อ 3.1: คงได้"
    if t.startswith("Feature ") and "URS" not in t and "ระบบ" in t:
        return f"หัวข้อ Feature ใน 3.1: เทียบ {C2}"
    if t.startswith("URS-"):
        return f"ข้อความ URS นี้: อ่านคู่ {C2}+{C1}; แก้ให้ครอบคลุมกฎลูกค้า; อัปเดตตาราง 3.2 ให้สอด"
    if "3.2 การจับคู่ URS" in t:
        return f"หัวข้อ 3.2: เพิ่มประโยคว่าตารางถัดไปต้องสอด {C2} และรหัส SRS ใช้ trace"
    if "4.1" in t and "Use Case Diagram" in t:
        return f"บท 4.1: เทียบภาพกับ {C2}; แก้ภาพถ้า actor/ขอบเขตต่าง"
    if "ภาพที่" in t or "Figure" in t:
        return f"คำบรรยายใต้ภาพ: เทียบ {C2}; อัปเดตเลขภาพให้ต่อเนื่องหลังแทรกภาพ"
    if "Activity Diagram" in t:
        return f"AD: เทียบลำดับกับ {C2}; แก้ swimlane/step ให้ตรง DnD/Reason"
    if "4.2" in t and "Use Case Description" in t:
        return "หัวข้อ 4.2: คงได้"
    if t.startswith("UC-"):
        return f"หัวข้อ UC: แก้ตารางถัดไป (หัวข้อ|รายละเอียด) ให้สอด {C2} ทุกฟิลด์ UC template"
    if "4.3" in t and "Activity Diagrams" in t:
        return "หัวข้อ 4.3: คงได้"
    if "ตารางข้อมูลการวางแผนงาน" in t or "Use Case Input" in t:
        return f"คำนำตารางฟิลด์: เทียบ {C2} — ตารางถัดไป (Table 11)"
    if "ตารางข้อมูลการเบิกอะไหล่" in t:
        return f"คำนำตารางฟิลด์: เทียบ {C2} — ตารางถัดไป (Table 13)"
    if "ตารางข้อมูลรายงาน" in t:
        return f"คำนำตารางรายงาน: เทียบ {C4}+{C2} — ตารางถัดไป (Table 14)"
    if "บทที่ 5" in t or "เอกสารอ้างอิง" in t:
        return f"เริ่มบทที่ 5: เพิ่มหัวข้อย่อยอ้างอิงลูกค้า {SET} พร้อมวันที่รับและเวอร์ชันไฟล์"
    if "Pepsi-Cola" in t or "SAP PM Module" in t or "Infrastructure Policy" in t:
        return "อ้างอิงภายในโรงงาน: คงได้; ตรวจชื่อเอกสารยังใช้งาน"
    if "[2]" in t or "29148" in t or "React.js" in t or "Docker Documentation" in t or "Tailscale Security" in t:
        return "อ้างอิงมาตรฐานสากล/เทคโนโลยี: คงได้; อัปเดตลิงก์เวอร์ชัน"
    if "Wikipedia" in t or "Telegram Bot API" in t:
        return "อ้างอิมออนไลน์: คงได้; ตรวจ URL"
    if "ภาคผนวก" in t and "Appendix" in t:
        return "หัวภาคผนวกรวม: คงได้"
    if "ภาคผนวก ก" in t:
        return "หัวภาคผนวก ก: คงได้"
    if "ภาคผนวก ข" in t:
        return f"หัวภาคผนวก ข: เทียบ RC กับ {C2}+{C1}"
    if t.startswith("RC0"):
        return f"รายการ RC: ต้องตรงรหัสและความหมายใน {C2} (รวม 05 อื่น ๆ ถ้ามี)"
    if "ภาคผนวก ค" in t:
        return "หัวภาคผนวก ค: คงได้"
    if "Tailscale VPN" in t or "Docker Ports" in t or t.startswith("Port "):
        return "รายละเอียดพอร์ต/ VPN: เทียบ INSTALL_SOP + รูป `from customer/server/`"

    if len(t) < 40 and ("โดย" in t or "ที่ปรึกษา" in t or "____" in t):
        return "ข้อมูลผู้จัดทำ/ลายเซ็น: ตรวจความเป็นปัจจุบัน"
    if re.match(r"^4\.2\.\d", t) and "โมดูล" in t:
        return f"หัวข้อย่อยบท 4.2: ตรวจว่า UC ด้านล่างสอด {C2}; ไม่ต้องแก้ถ้าเป็นชื่อโมดูลอย่างเดียว"
    if re.match(r"^4\.1\.\d", t) and "โมดูล" in t:
        return f"หัวข้อย่อยบท 4.1: เทียบภาพ/คำอธิบายกับ {C2}"
    return f"ย่อหน้าทั่วไป: อ่านคู่ {SET} — แก้ถ้าขัด {C2}; ถ้าเป็นคำนำไม่จำเป็นต้องเปลี่ยน"


def main() -> None:
    doc = Document(str(DOC))
    body = doc.element.body
    tbl_idx = -1
    rows_out: list[tuple[int, str, str, str]] = []
    seq = 0
    for child in body:
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            t = p.text.strip()
            if not t:
                continue
            seq += 1
            rows_out.append(
                (seq, "P", t[:280].replace("|", "/").replace("\n", " "), action_paragraph(t))
            )
        elif child.tag == qn("w:tbl"):
            tbl_idx += 1
            if tbl_idx < len(doc.tables):
                t = doc.tables[tbl_idx]
                r0 = " / ".join(c.text.strip().replace("\n", " ")[:45] for c in t.rows[0].cells)
                desc = f"Table{tbl_idx} {len(t.rows)}×{len(t.columns)}: {r0[:180]}"
            else:
                desc = f"Table{tbl_idx} (?)"
            seq += 1
            rows_out.append((seq, "T", desc, action_table(tbl_idx, desc)))

    lines = [
        "# รายการปรับแก้ SRS Pepsi — เรียงตามลำดับในเอกสาร Word",
        "",
        "ไฟล์อ้างอิง: `docs/Software Requirement Specification Pepsi Cola PM Project.docx`",
        "",
        "ถ้าต้องการแบบ **หัวข้อ → ปรับตรงไหน** (ไม่ไล่ทีละบรรทัด) ใช้ [`SRS_PEPSI_DOCX_REVISION_BY_HEADING.md`](SRS_PEPSI_DOCX_REVISION_BY_HEADING.md)",
        "",
        "**คอลัมน์ `#`:** ลำดับของย่อหน้า (`P`) และตาราง (`T`) ตามการปรากฏใน body ของไฟล์ Word (จากต้นจบเล่ม)",
        "",
        "**แหล่งความต้องการหลัก (ล็อกแบบ B):**",
        "",
        f"1. {C1}",
        f"2. {C2}",
        f"3. {C3}",
        f"4. {C4}",
        "",
        "**เอกสารช่วยใน repo:** [`PROJECT_PLAN.md`](PROJECT_PLAN.md), [`SRS_PEPSI_DOCX_REVISION_GUIDE.md`](SRS_PEPSI_DOCX_REVISION_GUIDE.md), [`SAP_DATA_IMPORT_EXPORT_COLUMNS.md`](SAP_DATA_IMPORT_EXPORT_COLUMNS.md), [`DATABASE_DESIGN_DRAFT.md`](DATABASE_DESIGN_DRAFT.md)",
        "",
        "สร้างซ้ำไฟล์นี้ได้ด้วย: `python scripts/build_srs_pepsi_item_revision_md.py`",
        "",
        "| # | ประเภท | เนื้อหาในเล่ม (ย่อ) | การปรับใน Word (ละเอียด) |",
        "|---:|:---|:---|:---|",
    ]
    for seq, kind, snippet, act in rows_out:
        esc = snippet.replace("|", "/")
        lines.append(f"| {seq} | {kind} | {esc} | {act} |")

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote", OUT, "rows", len(rows_out))


if __name__ == "__main__":
    main()
